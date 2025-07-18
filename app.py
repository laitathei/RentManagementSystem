import streamlit as st
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from decimal import Decimal, ROUND_HALF_UP
import pandas as pd
from datetime import datetime
import re
from io import BytesIO
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
import pytz

# ────────────────── 🔒 密碼登入驗證 ──────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
    st.session_state.user_name = ""

# 字體控制功能
if "font_size" not in st.session_state:
    st.session_state.font_size = 16
col_font1, col_font2 = st.columns([1, 1])
with col_font1:
    if st.button("🔍 放大字體"):
        st.session_state.font_size = min(st.session_state.font_size + 2, 32)
with col_font2:
    if st.button("🔎 縮小字體"):
        st.session_state.font_size = max(st.session_state.font_size - 2, 6)
st.markdown(f"""
    <style>
    html, body, [class*="css"]  {{
        font-size: {st.session_state.font_size}px !important;
    }}
    .stDataFrame div[data-testid="stDataFrame"] .ag-cell,
    .stDataEditor div[data-testid="stDataEditor"] .ag-cell {{
        font-size: {st.session_state.font_size}px !important;
    }}
    </style>
""", unsafe_allow_html=True)
# 主選單開始

if not st.session_state.authenticated:
    st.title("🏠 公司管理系統")
    st.subheader("🔐 請輸入密碼登入")

    with st.form("login_form"):
        pw = st.text_input("密碼", type="password")
        login_btn = st.form_submit_button("🔓 登入")
        if login_btn:
            pw2user = {v: k for k, v in st.secrets["USERS"].items()}   # 反轉成 {密碼:名字}
            if pw in pw2user:
                st.session_state.authenticated = True
                st.session_state.user_name = pw2user[pw]               # 記下誰登入
                st.success(f"✅ 歡迎 {st.session_state.user_name}！")
                st.rerun()
            else:
                st.error("❌ 密碼錯誤，請再試一次。")
    st.stop()  # ❗停止頁面，防止其他內容顯示

# ────────────────── Google Sheets 認證 ──────────────────
scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
# creds_dict = json.loads(st.secrets["GOOGLE_SERVICE_ACCOUNT_JSON"])
creds_dict = dict(st.secrets["GOOGLE_SERVICE_ACCOUNT_JSON"])
creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
client = gspread.authorize(creds)

sheet_id = "1k4EPnA9cLXkaDWpJ7EJwIwKpXWoth9mOMlHbNsFs644"
sheet_tenants   = client.open_by_key(sheet_id).worksheet("租客資料")  # 租客資料表
sheet_rentflow  = client.open_by_key(sheet_id).worksheet("租金流程")  # 租金流程表
sheet_listings  = client.open_by_key(sheet_id).worksheet("租賃盤源")  # 租金流程表

# ────────────────── 版頭 & 功能選單 ──────────────────
st.title("🏠 公司管理系統")
main_mode = st.radio("📂 功能類別", ["👥 租客資料管理", "📆 租金處理進度", "🏢 租賃盤源管理"], horizontal=True)

tenant_data   = sheet_tenants.get_all_records()
tenant_df     = pd.DataFrame(tenant_data)
tenant_df["sheet_order"] = tenant_df.index + 2          # 2,3,4,...
tenant_df["key"] = tenant_df["租客姓名"] + "｜" + tenant_df["單位地址"]
order_map = dict(zip(tenant_df["key"], tenant_df["sheet_order"]))
rentflow_data = sheet_rentflow.get_all_records()
rentflow_df   = pd.DataFrame(rentflow_data)
listing_data  = sheet_listings.get_all_records()
listing_df    = pd.DataFrame(listing_data)
for col in ["租客姓名", "單位地址", "租客電話", "固定水費", "固定電費", "每度水費", "每度電費", "起始水錶度數", "起始電錶度數"]:
    if col in tenant_df.columns:
        tenant_df[col] = tenant_df[col].astype(str)
for col in ["本月水錶度數", "上月水錶度數", "本月水費", "本月電錶度數", "上月電錶度數", "本月電費"]:
    if col in rentflow_df.columns:
        rentflow_df[col] = rentflow_df[col].astype(str)
for col in ["收租金額", "過戶金額"]:
    if col in rentflow_df.columns:
        rentflow_df[col] = pd.to_numeric(rentflow_df[col], errors="coerce")
for col in ["建築面積", "租金要求"]:
    if col in listing_df.columns:
        listing_df[col] = pd.to_numeric(listing_df[col], errors="coerce")
for col in ["最多人數限制", "業主電話"]:
    if col in listing_df.columns:
        listing_df[col] = listing_df[col].astype(str)

if main_mode == "👥 租客資料管理":
    # 讀取資料
    st.subheader("📋 租客資料")
    st.data_editor(tenant_df.set_index(pd.RangeIndex(start=1, stop=len(tenant_df)+1)), use_container_width=True, disabled=True)
    sub_mode = st.radio("📋 租客操作選項", ["➕ 新增租客資料", "✏️ 更改租客資料", "🗑️ 刪除租客資料"], horizontal=True)

    # ────────────────── 新增 ──────────────────
    if sub_mode == "➕ 新增租客資料":
        st.subheader("➕ 新增租客資料")
        water_mode = st.radio("💧 水費收費方式", ["每度計算", "固定金額", "不代收"], horizontal=True)
        electric_mode = st.radio("⚡ 電費收費方式", ["每度計算", "固定金額", "不代收"], horizontal=True)
        with st.form("add_tenant_form"):
            name = st.text_input("租客姓名")
            phone = st.text_input("電話")
            address = st.text_input("單位地址")
            rent = st.number_input("每月固定租金", min_value=0.0)

            # 水費
            water_box = st.empty()
            if water_mode == "每度計算":
                water_fee = st.number_input("每度水費", min_value=0.0, key="water_per_unit_add")
                fix_water_fee = "N/A"
            elif water_mode == "固定金額":
                fix_water_fee = st.number_input("固定水費金額", min_value=0.0, key="water_fixed_add")
                water_fee = "N/A"
            else:
                water_fee = fix_water_fee = "N/A"

            # 電費
            electric_box = st.empty()
            if electric_mode == "每度計算":
                electric_fee = st.number_input("每度電費", min_value=0.0, key="electric_per_unit_add")
                fix_electric_fee = "N/A"
            elif electric_mode == "固定金額":
                fix_electric_fee = st.number_input("固定電費金額", min_value=0.0, key="electric_fixed_add")
                electric_fee = "N/A"
            else:
                electric_fee = fix_electric_fee = "N/A"
            
            init_water_units = st.number_input("起始水錶度數", min_value=0.0, key="init_water_units")
            init_elec_units = st.number_input("起始電錶度數", min_value=0.0, key="init_elec_units")

            language = st.selectbox("通訊語言", ["中文", "英文"])
            management_fee = st.number_input("收租費", min_value=0.0, value=0.0)
            cutoff_day = st.selectbox("截數日（每月）", list(range(1, 32)))
            lease_type = st.selectbox("租約狀態", ["新租", "續租"], index=0)
            lease_start = st.date_input("租約開始日", key="lease_start", value=pd.Timestamp.now().date())
            lease_end   = st.date_input("租約結束日", key="lease_end", value=pd.Timestamp.now().date() + pd.DateOffset(years=2))

            if st.form_submit_button("✅ 新增"):
                tz_hk = pytz.timezone("Asia/Hong_Kong")
                ts = datetime.now(tz_hk).strftime("%Y-%m-%d %H:%M:%S")
                who = st.session_state.get("user_name", "unknown")
                exists = rentflow_df[
                    (tenant_df["租客姓名"] == name.strip()) &
                    (tenant_df["單位地址"] == address.strip())
                ]
                if not exists.empty:
                    st.warning("⚠️ 已存在相同租客姓名與單位地址的紀錄，請確認是否重覆輸入。")
                    st.stop()
                else:
                    new_row = [name, phone, address, rent, fix_water_fee, fix_electric_fee, water_fee, electric_fee, init_water_units, init_elec_units,
                            cutoff_day, language, management_fee, lease_type, str(lease_start), str(lease_end), ts, who]
                    sheet_tenants.append_row(new_row)
                    st.success(f"✅ 已新增租客：{name}")
                    st.rerun()

    # ────────────────── 更改 ──────────────────
    elif sub_mode == "✏️ 更改租客資料":
        st.subheader("✏️ 更改租客資料")
        if tenant_df.empty:
            st.info("目前沒有資料可修改。")
        else:
            selector = tenant_df["租客姓名"] + "｜" + tenant_df["單位地址"]
            choice = st.selectbox("選擇欲修改的租客", selector)
            idx = selector.tolist().index(choice)
            sheet_row = idx + 2  # Sheet 2-based
            row = tenant_df.iloc[idx]

            water_mode_options    = ["每度計算", "固定金額", "不代收"]
            electric_mode_options = ["每度計算", "固定金額", "不代收"]

            if str(row["每度水費"]).upper() != "N/A" and str(row["每度水費"]) != "":
                water_mode_idx   = water_mode_options.index("每度計算")
            elif str(row["固定水費"]).upper() != "N/A" and str(row["固定水費"]) != "":
                water_mode_idx   = water_mode_options.index("固定金額")
            else:
                water_mode_idx   = water_mode_options.index("不代收")


            if str(row["每度電費"]).upper() != "N/A" and str(row["每度電費"]) != "":
                electric_mode_idx   = electric_mode_options.index("每度計算")
            elif str(row["固定電費"]).upper() != "N/A" and str(row["固定電費"]) != "":
                electric_mode_idx   = electric_mode_options.index("固定金額")
            else:
                electric_mode_idx   = electric_mode_options.index("不代收")

            water_mode = st.radio("💧 水費收費方式", water_mode_options, index=water_mode_idx, horizontal=True)
            electric_mode = st.radio("⚡ 電費收費方式", electric_mode_options, index=electric_mode_idx, horizontal=True)
        
            with st.form("edit_tenant_form"):
                name    = st.text_input("租客姓名",  value=row["租客姓名"])
                phone   = st.text_input("電話",      value=str(row["租客電話"]))
                address = st.text_input("單位地址",  value=row["單位地址"])
                rent    = st.number_input("每月固定租金", value=float(row["每月固定租金"]))

                # 水費
                water_box = st.empty()
                if water_mode == "每度計算":
                    water_fee = st.number_input("每度水費", min_value=0.0, 
                                                value=float(row["每度水費"]) if str(row["每度水費"]).replace('.', '', 1).isdigit() else 0.0,
                                                key="water_per_unit_add")
                    fix_water_fee = "N/A"
                elif water_mode == "固定金額":
                    fix_water_fee = st.number_input("固定水費金額", min_value=0.0, 
                                                    value=float(row["固定水費"]) if str(row["固定水費"]).replace('.', '', 1).isdigit() else 0.0,
                                                    key="water_fixed_add")
                    water_fee = "N/A"
                else:
                    water_fee = fix_water_fee = "N/A"
            
                # 電費
                electric_box = st.empty()
                if electric_mode == "每度計算":
                    electric_fee = st.number_input("每度電費", min_value=0.0, 
                                                   value=float(row["每度電費"]) if str(row["每度電費"]).replace('.', '', 1).isdigit() else 0.0,
                                                   key="electric_per_unit_add")
                    fix_electric_fee = "N/A"
                elif electric_mode == "固定金額":
                    fix_electric_fee = st.number_input("固定電費金額", min_value=0.0, 
                                                       value=float(row["固定電費"]) if str(row["固定電費"]).replace('.', '', 1).isdigit() else 0.0,
                                                       key="electric_fixed_add")
                    electric_fee = "N/A"
                else:
                    electric_fee = fix_electric_fee = "N/A"
    
                init_water_units = st.number_input("起始水錶度數", min_value=0.0, value=float(row["起始水錶度數"]), key="init_water_units")
                init_elec_units = st.number_input("起始電錶度數", min_value=0.0, value=float(row["起始電錶度數"]), key="init_elec_units")

                language = st.selectbox("通訊語言", ["中文", "英文"], index=0 if row["通訊語言"]=="中文" else 1)
                management_fee = st.number_input("收租費", min_value=0.0, value=float(row["收租費"]))
                cutoff_day = st.selectbox("截數日（每月）", list(range(1, 32)), index=int(row["截數日"])-1)
                lease_type = st.selectbox("租約狀態", ["新租", "續租"], index = 0 if row["租約狀態"] != "續租" else 1)
                lease_start = st.date_input("租約開始日", value=pd.to_datetime(row["租約開始日"]) if "租約開始日" in row and row["租約開始日"] else pd.Timestamp.now().date())
                lease_end   = st.date_input("租約結束日", value=pd.to_datetime(row["租約結束日"]) if "租約結束日" in row and row["租約結束日"] else pd.Timestamp.now().date() + pd.DateOffset(years=2))

                if st.form_submit_button("💾 儲存修改"):
                    tz_hk = pytz.timezone("Asia/Hong_Kong")
                    ts = datetime.now(tz_hk).strftime("%Y-%m-%d %H:%M:%S")
                    who = st.session_state.get("user_name", "unknown")
                    new_row = [name, phone, address, rent, 
                               fix_water_fee, fix_electric_fee, 
                               water_fee, electric_fee,
                               cutoff_day, language, management_fee, lease_type,
                               str(lease_start), str(lease_end),
                               ts, who]
                    sheet_tenants.update(f"A{sheet_row}:O{sheet_row}", [new_row])
                    st.success("✅ 已更新！")
                    st.rerun()

    # ────────────────── 刪除 ──────────────────
    elif sub_mode == "🗑️ 刪除租客資料":
        st.subheader("🗑️ 刪除租客資料")
        if tenant_df.empty:
            st.info("目前沒有資料可刪除。")
        else:
            selector = tenant_df["租客姓名"] + "｜" + tenant_df["單位地址"]
            choice = st.selectbox("選擇欲刪除的租客", selector)
            idx = selector.tolist().index(choice)
            sheet_row = idx + 2  # Sheet 2-based

            if st.button("⚠️ 確認刪除"):
                sheet_tenants.delete_rows(sheet_row)
                st.warning(f"已刪除：{choice}")
                st.rerun()

elif main_mode == "📆 租金處理進度":
    st.markdown("### 🔍 指定月份查詢")
    now = datetime.now()
    all_years  = sorted(set(rentflow_df["年度"].unique().tolist() + [now.year]), reverse=True)
    all_months = sorted(set(rentflow_df["月份"].unique().tolist() + [now.month]))
    selected_year  = st.selectbox("選擇年份", all_years, index=0)
    selected_month = st.selectbox("選擇月份", all_months, index=all_months.index(now.month))
    month_start = pd.Timestamp(selected_year, selected_month, 1)
    month_end = pd.Timestamp(selected_year, selected_month, 1) + pd.offsets.MonthEnd(0)
    filtered_df = rentflow_df[
        (rentflow_df["年度"] == selected_year) &
        (rentflow_df["月份"] == selected_month)
    ]

    # ────── ❷ 將租約日期欄位轉為 datetime，方便比對 ──────
    tenant_df["租約開始日"] = pd.to_datetime(tenant_df["租約開始日"], errors="coerce")
    tenant_df["租約結束日"] = pd.to_datetime(tenant_df["租約結束日"], errors="coerce")

    # ────── ❸ 只挑出「本月需要交租」的租客  ──────
    # 續租 = 一律要收
    cond_renew = (
        (tenant_df["租約狀態"] == "續租") &
        (tenant_df["租約開始日"] <= month_end)
    )
    # 新租 = 起租日在本月 1 號「之前」才要收(即首租期由下一月開始)
    cond_new   = (
        (tenant_df["租約狀態"] != "續租") &            # 空白或「新租」
        (tenant_df["租約開始日"] < month_start)        # 嚴格 < 本月 1 日
    )

    active_df = tenant_df[cond_renew | cond_new].copy()

    st.markdown(f"### 📋 {selected_year} 年 {selected_month} 月租金流程")
    active_df["key"]   = active_df["租客姓名"] + "｜" + active_df["單位地址"].astype(str)
    filtered_df["key"] = filtered_df["租客姓名"] + "｜" + filtered_df["單位地址"].astype(str)
    
    calculated_df = filtered_df[filtered_df["已計算水電"].astype(str).str.upper() == "TRUE"]
    calculated_rooms = len(calculated_df)
    calculated_keys  = set(calculated_df["key"])
    # ① 未計算 = 月內「應收」但 key 不在 calc_keys
    uncalculated_df = active_df[~active_df["key"].isin(calculated_keys)]
    uncalculated_df = (uncalculated_df
        .assign(索引=lambda d: d["key"].map(order_map))
        .sort_values("索引")
        .set_index("索引")          # <<── index 就是 9‧10‧12…
    )

    paid_df   = filtered_df[filtered_df["已收取租金"].astype(str).str.upper() == "TRUE"]
    paid_rooms = len(paid_df)                         # ← 行數就是房間數
    paid_keys  = set(paid_df["key"])                  # ← 用來做未交租比對
    # ② 未收租  = 已經計算 (key 在 calc_keys) 但還沒 paid
    unpaid_df = active_df[active_df["key"].isin(calculated_keys) & ~active_df["key"].isin(paid_keys)]
    unpaid_df = (unpaid_df
        .assign(索引=lambda d: d["key"].map(order_map))
        .sort_values("索引")
        .set_index("索引")          # <<── index 就是 9‧10‧12…
    )
    unpaid_rooms = len(unpaid_df)           # 未交租房間數
    
    deposit_df = filtered_df[filtered_df["已存入租金"].astype(str).str.upper() == "TRUE"]
    deposited_rooms = len(deposit_df)
    deposit_keys = set(deposit_df["key"])
    # ③ 未入帳  = 已收租且 key 在 paid_keys，但不在 dep_keys
    undeposited_df = filtered_df[(filtered_df["key"].isin(paid_keys)) & (~filtered_df["key"].isin(deposit_keys))]
    undeposited_df = (undeposited_df
        .assign(索引=lambda d: d["key"].map(order_map))
        .sort_values("索引")
        .set_index("索引")          # <<── index 就是 9‧10‧12…
    )
    undeposited_rooms = len(undeposited_df)
    total_rooms  = len(active_df)                     # 全部房間

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("📋 總租客數", total_rooms)
    col2.metric("🧮 已計算水電", calculated_rooms)
    col3.metric("✅ 已交租", paid_rooms)
    col4.metric("🏦 已入帳", deposited_rooms)
    st.data_editor(filtered_df.drop(columns=["key"]).set_index(pd.RangeIndex(start=1, stop=len(filtered_df.drop(columns=["key"]))+1)), use_container_width=True, disabled=True)

    if uncalculated_df.empty: # 已計算水電
        st.success(f"🥳 所有 {selected_year} 年 {selected_month} 月租客都已完成水電計算")
    else: # 尚未計算水電
        st.markdown("### 🧮 尚未計算水電名單")
        cols = [c for c in ["租客姓名", "租客電話", "單位地址"] if c in uncalculated_df.columns]
        st.data_editor(uncalculated_df[cols], use_container_width=True, disabled=True)

    if unpaid_df.empty and uncalculated_df.empty: # 已計算水電和已收租
        st.success(f"🥳 所有 {selected_year} 年 {selected_month} 月租客都已完成收租")
    else: # 已計算水電但未收租
        st.markdown("### ❌ 已計算水電但尚未收租名單")
        tmp = filtered_df[["key", "水電金額"]].copy()
        tmp["水電金額"] = pd.to_numeric(tmp["水電金額"], errors="coerce").fillna(0)
        unpaid_view = (unpaid_df.merge(tmp, on="key", how="left"))
        unpaid_view["應付金額"] = (pd.to_numeric(unpaid_view["每月固定租金"], errors="coerce").fillna(0) + unpaid_view["水電金額"])
        cols = [c for c in ["租客姓名", "租客電話", "單位地址", "應付金額"] if c in unpaid_df.columns]
        st.data_editor(unpaid_df[cols], use_container_width=True, disabled=True)

    if unpaid_df.empty and uncalculated_df.empty and undeposited_df.empty: # 已計算水電和已收租和已過戶
        st.success(f"🥳 所有 {selected_year} 年 {selected_month} 月租客都已完成過戶")
    else: # 已收租但未入帳
        st.markdown("### 🏦 已收租但尚未過數名單")
        cols = [c for c in ["租客姓名", "租客電話", "單位地址", "收租金額", "收取租金日期"] if c in undeposited_df.columns]
        st.data_editor(undeposited_df[cols], use_container_width=True, disabled=True)

    sub_mode = st.radio("🧾 租金紀錄操作", ["➕ 新增租金紀錄", "✏️ 更改租金紀錄", "🗑️ 刪除租金紀錄", "📄 產生業主收據"], horizontal=True)
    if sub_mode == "➕ 新增租金紀錄":
        st.subheader("➕ 新增租金紀錄")

        pending_df = active_df[~active_df["key"].isin(deposit_keys)]

        if pending_df.empty:
            st.info("🥳 所有租客都已繳交該月份租金，無需新增紀錄。")
            st.stop()

        selector = pending_df["租客姓名"] + "｜" + pending_df["單位地址"]
        sel_opt = st.selectbox("租客", selector)

        idx = selector.tolist().index(sel_opt)
        default_phone = str(pending_df.iloc[idx]["租客電話"]).lstrip("'").strip()
        name = sel_opt.split("｜")[0]
        address = pending_df.iloc[idx]["單位地址"]
        default_rent = float(pending_df.iloc[idx]["每月固定租金"])

        calculate_done  = st.checkbox("🧮 已計算費用", key="calculate_done_out")
        receive_done  = st.checkbox("✅ 已收租", key="receive_done_out")
        deposit_done  = st.checkbox("🏦 已入帳", key="deposit_done_out")

        with st.form("add_rentflow_form"):
            phone = st.text_input("租客電話", value=default_phone, disabled=True)

            year = st.number_input("年度", min_value=2000, max_value=2100, value=pd.Timestamp.now().year)
            month = st.selectbox("月份", list(range(1, 13)), index=pd.Timestamp.now().month - 1)

            trow = pending_df.iloc[idx]                        # 取得該租客在 tenant_df 的資料
            hist_df = rentflow_df[
                (rentflow_df["租客姓名"] == name) &
                (rentflow_df["單位地址"] == address) &
                (
                    (rentflow_df["年度"] < year) |
                    ((rentflow_df["年度"] == year) & (rentflow_df["月份"] < month))
                )
            ]

            # ➋ 取出最近一筆（年度、月份都最大的那行）
            to_number_or_na = lambda v: float(v) if isinstance(v, (int, float)) or (isinstance(v, str) and v.replace('.', '', 1).isdigit()) else v
            if not hist_df.empty:
                # 先按 年度、月份 由大到小排序，再取第一筆
                prev_row = hist_df.sort_values(["年度", "月份"], ascending=False).iloc[0]
                prev_water_units = to_number_or_na(prev_row["本月水錶度數"])
                prev_elec_units  = to_number_or_na(prev_row["本月電錶度數"])
            else:
                # 找不到任何舊紀錄，就用租客資料的「起始錶度數」
                prev_water_units = to_number_or_na(trow["起始水錶度數"])
                prev_elec_units  = to_number_or_na(trow["起始電錶度數"])

            if str(trow["每度水費"]).upper() != "N/A" and str(trow["每度水費"]) != "":
                water_mode = "per_unit"          # 按度數計費
            elif str(trow["固定水費"]).upper() != "N/A" and str(trow["固定水費"]) != "":
                water_mode = "fixed"             # 固定金額
            else:
                water_mode = "none"              # 不代收

            if str(trow["每度電費"]).upper() != "N/A" and str(trow["每度電費"]) != "":
                elec_mode = "per_unit"          # 按度數計費
            elif str(trow["固定電費"]).upper() != "N/A" and str(trow["固定電費"]) != "":
                elec_mode = "fixed"             # 固定金額
            else:
                elec_mode = "none"              # 不代收

            sig_val = lambda v: "N/A" if (v is None or (isinstance(v, str) and v.strip() == "")) else str(v)
            if calculate_done:
                if water_mode == "per_unit":
                    curr_water_units = st.number_input("💧 本月水錶度數", min_value=0.0, step=0.1, value=st.session_state.get("curr_water_units", 0.0), key="curr_water_units")
                else:
                    curr_water_units = prev_water_units
                if elec_mode == "per_unit":
                    curr_elec_units  = st.number_input("⚡ 本月電錶度數", min_value=0.0, step=0.1, value=st.session_state.get("curr_elec_units", 0.0), key="curr_elec_units")
                else:
                    curr_elec_units  = prev_elec_units
                calculate_date = st.date_input("📅 計算日期", value=pd.Timestamp.now().date(), key="calculated_date_in")

                if st.form_submit_button("🔢 計算"):
                    # ② 計算水費
                    if water_mode == "per_unit":
                        water_units = max(0, round(float(curr_water_units) - float(prev_water_units)))
                        value = Decimal(water_units) * Decimal(str(trow["每度水費"]))
                        water_fee = int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
                    elif water_mode == "fixed":
                        water_fee = float(trow["固定水費"])
                        water_units = "N/A"
                    else:
                        water_fee = "N/A"
                        water_units = "N/A"

                    # ③ 計算電費
                    if elec_mode == "per_unit":
                        elec_units  = max(0, round(float(curr_elec_units)  - float(prev_elec_units)))
                        value = Decimal(elec_units) * Decimal(str(trow["每度電費"]))
                        elec_fee = int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
                    elif elec_mode == "fixed":
                        elec_fee = float(trow["固定電費"])
                        elec_units = "N/A"
                    else:
                        elec_fee = "N/A"
                        elec_units = "N/A"

                    to_float_safe = lambda v: float(v) if isinstance(v, (int, float)) or (isinstance(v, str) and v.replace('.', '', 1).isdigit()) else 0
                    calculate_amt = default_rent + to_float_safe(water_fee) + to_float_safe(elec_fee)
                    water_elec_fee = to_float_safe(water_fee) + to_float_safe(elec_fee)

                    # ⬇︎ 把結果暫存，供後面「新增」使用
                    st.session_state["rent_calc"] = {
                        "water_units": water_units,
                        "elec_units": elec_units,
                        "water_fee": water_fee,
                        "elec_fee": elec_fee,
                        "water_elec_fee": water_elec_fee,
                        "calculate_amt": calculate_amt,
                        "calculate_date": calculate_date,
                        "inputs": (year, month, sig_val(curr_water_units), sig_val(curr_elec_units))
                    }

                    if "rent_calc" in st.session_state:
                        rc = st.session_state["rent_calc"]

                        # ➊ 水錶資訊一行
                        if water_mode == "per_unit":
                            col1, col2, col3 = st.columns(3)
                            col1.info(f"💧 本月水錶: {float(curr_water_units)}")
                            col2.info(f"💧 上月水錶: {float(prev_water_units)}")
                            col3.info(f"💧 每度水費: HK$ {float(trow['每度水費'])}")

                        # ➋ 電錶資訊一行
                        if elec_mode == "per_unit":
                            col4, col5, col6 = st.columns(3)
                            col4.info(f"⚡ 本月電錶: {float(curr_elec_units)}")
                            col5.info(f"⚡ 上月電錶: {float(prev_elec_units)}")
                            col6.info(f"⚡ 每度電費: HK$ {float(trow['每度電費'])}")

                        # ➌ 金額一行（水費／電費／租金）
                        col7, col8, col9 = st.columns(3)
                        col7.info(f"💧 水費: HK$ {rc['water_fee']}")
                        col8.info(f"⚡ 電費: HK$ {rc['elec_fee']}")
                        col9.info(f"💰 租金: HK$ {default_rent}")

                        # ➍ 總金額一行
                        st.info(f"📘 合共: HK$ {rc['calculate_amt']}")
            else:
                water_fee = ""
                elec_fee = ""
                water_elec_fee = ""
                calculate_date = ""
                calculate_amt = ""
                st.session_state.pop("rent_calc", None)   # 取消勾選時清空

            rc = st.session_state.get("rent_calc", {})
            init_receive = rc.get("calculate_amt", default_rent)   # 月租 + 水電 OR 月租
            if receive_done:
                receive_date = st.date_input("📅 收租日期", value=pd.Timestamp.now().date(), key="receive_date_in")
                receive_amt  = st.number_input("💰 收租金額", min_value=0.0, value=init_receive, key="receive_amt")
            else:
                receive_date = ""
                receive_amt = ""
            if deposit_done:
                deposit_date = st.date_input("📅 過數日期", value=pd.Timestamp.now().date(), key="deposit_date_in")
                deposit_amt  = st.number_input("💰 過戶金額", min_value=0.0, value=init_receive, key="deposit_amt")
            else:
                deposit_date = ""
                deposit_amt = ""

            # ───── 小工具：把值標準化成可比對的字串 ─────
            norm = lambda v: "N/A" if v in (None, "", "N/A") else str(v)

            # ① 每次『租客 selector』改變時，把舊的輸入清掉
            if "last_selector" not in st.session_state or st.session_state.last_selector != sel_opt:
                for k in ("curr_water_units", "curr_elec_units", "rent_calc"):
                    st.session_state.pop(k, None)
                st.session_state.last_selector = sel_opt      # 記住這次選的人

            # ② 已計算費用 = True／False 時，也要同步清掉舊的計算
            if not calculate_done:
                for k in ("curr_water_units", "curr_elec_units", "rent_calc"):
                    st.session_state.pop(k, None)

            # ③ 取得目前輸入（沒有就給 "N/A"）
            cur_w = norm(st.session_state.get("curr_water_units", "N/A"))
            cur_e = norm(st.session_state.get("curr_elec_units", "N/A"))

            # ④ 檢查試算是否仍然有效
            rc = st.session_state.get("rent_calc")
            calc_ok = (
                calculate_done and
                rc and
                (year, month, norm(rc["inputs"][2]), norm(rc["inputs"][3])) ==
                (year, month, cur_w, cur_e)
            )

            if not calc_ok and calculate_done:
                st.warning("⚠️ 請先按『🔢 計算』計算金額，再儲存！")
                st.stop()

            if st.form_submit_button("✅ 新增", disabled = not calc_ok):
                water_units     = rc.get("water_units", "")
                prev_water_units = prev_water_units          # 仍沿用先前計算好的舊度數
                water_fee       = rc.get("water_fee", "")
                elec_units      = rc.get("elec_units", "")
                prev_elec_units = prev_elec_units
                elec_fee        = rc.get("elec_fee", "")
                water_elec_fee  = rc.get("water_elec_fee", "")
                calculate_amt   = rc.get("calculate_amt", "")
                calculate_date  = rc.get("calculate_date", "")

                tz_hk = pytz.timezone("Asia/Hong_Kong")
                ts = datetime.now(tz_hk).strftime("%Y-%m-%d %H:%M:%S")
                who = st.session_state.get("user_name", "unknown")
                exists = rentflow_df[
                    (rentflow_df["租客姓名"] == name) &
                    (rentflow_df["單位地址"] == address) &
                    (rentflow_df["年度"] == year) &
                    (rentflow_df["月份"] == month)
                ]
                if not exists.empty:
                    st.warning(f"⚠️ 此租客{selected_year} 年 {selected_month} 月的租金流程紀錄已存在！")
                    st.stop()
                else:
                    row = [
                        phone, name, address, year, month,
                        str(calculate_date) if calculate_done else "",
                        calculate_done,
                        water_elec_fee  if calculate_done  else "",
                        str(receive_date) if receive_done else "",
                        receive_done,
                        receive_amt  if receive_done  else "",
                        str(deposit_date) if deposit_done else "",
                        deposit_done,
                        deposit_amt  if deposit_done else "",
                        curr_water_units, prev_water_units, water_fee,
                        curr_elec_units, prev_elec_units, elec_fee,
                        calculate_amt,
                        ts, who
                    ]
                    sheet_rentflow.append_row(row, value_input_option="RAW")
                    st.success("✅ 已成功新增租金紀錄")
                    for k in ("curr_water_units", "curr_elec_units"):
                        st.session_state.pop(k, None)   # 刪掉就會回到 default value
                    st.rerun()

    elif sub_mode == "✏️ 更改租金紀錄":
        st.subheader("✏️ 更改租金紀錄")
        if filtered_df.empty:
            st.info(f"目前沒有 {selected_year} 年 {selected_month} 月的紀錄可修改")
        else:
            filtered_df["選項"] = (
                filtered_df["租客姓名"] + "｜" +
                filtered_df["單位地址"] + "｜" +
                filtered_df["年度"].astype(str) + "-" + filtered_df["月份"].astype(str).str.zfill(2)
            )
            rentflow_df["選項"] = (
                filtered_df["租客姓名"] + "｜" +
                filtered_df["單位地址"] + "｜" +
                filtered_df["年度"].astype(str) + "-" + filtered_df["月份"].astype(str).str.zfill(2)
            )
            sel_opt = st.selectbox("選擇要修改的紀錄", filtered_df["選項"].tolist())
            idx = rentflow_df[rentflow_df["選項"] == sel_opt].index[0]
            row_data = rentflow_df.loc[idx]
            gs_row = idx + 2  # Google Sheets 的列數（從第2列開始）

            name = row_data["租客姓名"]
            address = row_data["單位地址"]
            trow = tenant_df[(tenant_df["租客姓名"] == name) & (tenant_df["單位地址"] == address)].iloc[0]
            default_rent = float(trow["每月固定租金"])
            if str(trow["每度水費"]).upper() != "N/A" and str(trow["每度水費"]) != "":
                water_mode = "per_unit"          # 按度數計費
            elif str(trow["固定水費"]).upper() != "N/A" and str(trow["固定水費"]) != "":
                water_mode = "fixed"             # 固定金額
            else:
                water_mode = "none"              # 不代收

            if str(trow["每度電費"]).upper() != "N/A" and str(trow["每度電費"]) != "":
                elec_mode = "per_unit"          # 按度數計費
            elif str(trow["固定電費"]).upper() != "N/A" and str(trow["固定電費"]) != "":
                elec_mode = "fixed"             # 固定金額
            else:
                elec_mode = "none"              # 不代收

            calculate_done  = st.checkbox("🧮 已計算費用", value=str(row_data["已計算水電"]).upper() == "TRUE")
            receive_done = st.checkbox("✅ 已收租", value=str(row_data["已收取租金"]).upper() == "TRUE")
            deposit_done = st.checkbox("🏦 已入帳", value=str(row_data["已存入租金"]).upper() == "TRUE")
            
            hist_df = rentflow_df[
                (rentflow_df["租客姓名"] == name) &
                (rentflow_df["單位地址"] == address) &
                (
                    (rentflow_df["年度"] < selected_year) |
                    ((rentflow_df["年度"] == selected_year) & (rentflow_df["月份"] < selected_month))
                )
            ]

            # ➋ 取出最近一筆（年度、月份都最大的那行）
            to_number_or_na = lambda v: float(v) if isinstance(v, (int, float)) or (isinstance(v, str) and v.replace('.', '', 1).isdigit()) else v
            if not hist_df.empty:
                # 先按 年度、月份 由大到小排序，再取第一筆
                prev_row = hist_df.sort_values(["年度", "月份"], ascending=False).iloc[0]
                prev_water_units = to_number_or_na(prev_row["本月水錶度數"])
                prev_elec_units  = to_number_or_na(prev_row["本月電錶度數"])
            else:
                # 找不到任何舊紀錄，就用租客資料的「起始錶度數」
                prev_water_units = to_number_or_na(trow["起始水錶度數"])
                prev_elec_units  = to_number_or_na(trow["起始電錶度數"])

            with st.form("edit_rentflow_form"):
                sig_val = lambda v: "N/A" if (v is None or (isinstance(v, str) and v.strip() == "")) else str(v)
                if calculate_done:
                    if water_mode == "per_unit":
                        curr_water_units = st.number_input("💧 本月水錶度數", min_value=0.0, step=0.1, value=float(row_data["本月水錶度數"]), key="curr_water_units")
                        water_units = max(0, round(float(curr_water_units) - float(prev_water_units)))
                        value = Decimal(water_units) * Decimal(str(trow["每度水費"]))
                        water_fee = int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
                    elif water_mode == "fixed":
                        curr_water_units = "N/A"
                        prev_water_units = "N/A"
                        water_fee = float(trow["固定水費"])
                        water_units = "N/A"
                    else:
                        curr_water_units = "N/A"
                        prev_water_units = "N/A"
                        water_fee = "N/A"
                        water_units = "N/A"

                    if elec_mode == "per_unit":
                        curr_elec_units  = st.number_input("⚡ 本月電錶度數", min_value=0.0, step=0.1, value=float(row_data["本月電錶度數"]), key="curr_elec_units")
                        elec_units  = max(0, round(float(curr_elec_units)  - float(prev_elec_units)))
                        value = Decimal(elec_units) * Decimal(str(trow["每度電費"]))
                        elec_fee = int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
                    elif elec_mode == "fixed":
                        curr_elec_units = "N/A"
                        prev_elec_units = "N/A"
                        elec_fee = float(trow["固定電費"])
                        elec_units = "N/A"
                    else:
                        curr_elec_units = "N/A"
                        prev_elec_units = "N/A"
                        elec_fee = "N/A"
                        elec_units = "N/A"

                    to_float_safe = lambda v: float(v) if isinstance(v, (int, float)) or (isinstance(v, str) and v.replace('.', '', 1).isdigit()) else 0
                    calculate_amt = default_rent + to_float_safe(water_fee) + to_float_safe(elec_fee)
                    water_elec_fee = to_float_safe(water_fee) + to_float_safe(elec_fee)
                    
                    calculate_date = st.date_input("📅 計算日期", value=pd.to_datetime(row_data["計算日期"]).date() if row_data["計算日期"] else pd.Timestamp.now().date(), key="calculate_date_in")
                    if st.form_submit_button("🔢 計算"):
                        # ⬇︎ 把結果暫存，供後面「新增」使用
                        st.session_state["modify_calc"] = {
                            "water_units": water_units,
                            "elec_units": elec_units,
                            "water_fee": water_fee,
                            "elec_fee": elec_fee,
                            "water_elec_fee": water_elec_fee,
                            "calculate_amt": calculate_amt,
                            "calculate_date": calculate_date,
                            "inputs": (selected_year, selected_month, sig_val(curr_water_units), sig_val(curr_elec_units))
                        }

                    if "modify_calc" in st.session_state:
                        rc = st.session_state["modify_calc"]

                        # ➊ 水錶資訊一行
                        if water_mode == "per_unit":
                            col1, col2, col3 = st.columns(3)
                            col1.info(f"💧 本月水錶: {float(curr_water_units)}")
                            col2.info(f"💧 上月水錶: {float(prev_water_units)}")
                            col3.info(f"💧 每度水費: HK$ {float(trow['每度水費'])}")
                            
                        # ➋ 電錶資訊一行
                        if elec_mode == "per_unit":
                            col4, col5, col6 = st.columns(3)
                            col4.info(f"⚡ 本月電錶: {float(curr_elec_units)}")
                            col5.info(f"⚡ 上月電錶: {float(prev_elec_units)}")
                            col6.info(f"⚡ 每度電費: HK$ {float(trow['每度電費'])}")

                        # ➌ 金額一行（水費／電費／租金）
                        col7, col8, col9 = st.columns(3)
                        col7.info(f"💧 水費: HK$ {rc['water_fee']}")
                        col8.info(f"⚡ 電費: HK$ {rc['elec_fee']}")
                        col9.info(f"💰 租金: HK$ {default_rent}")

                        # ➍ 總金額一行
                        st.info(f"📘 合共: HK$ {rc['calculate_amt']}")
                else:
                    water_fee = ""
                    elec_fee = ""
                    water_elec_fee = ""
                    calculate_date = ""
                    calculate_amt = ""
                    st.session_state.pop("modify_calc", None)   # 取消勾選時清空

                rc = st.session_state.get("modify_calc", {})
                init_receive = rc.get("calculate_amt", default_rent)   # 月租 + 水電 OR 月租
                if receive_done:
                    receive_date = st.date_input("📅 收租日期", value=pd.to_datetime(row_data["收取租金日期"]).date() if row_data["收取租金日期"] else pd.Timestamp.now().date(), key="receive_date_in")
                    receive_amt  = st.number_input("💰 收租金額", min_value=0.0, value=init_receive, key="receive_amt")
                else:
                    receive_date = ""
                    receive_amt = ""
                if deposit_done:
                    deposit_date = st.date_input("📅 過數日期", value=pd.to_datetime(row_data["存入租金日期"]).date() if row_data["存入租金日期"] else pd.Timestamp.now().date(), key="deposit_date_in")
                    deposit_amt  = st.number_input("💰 過戶金額", min_value=0.0, value=init_receive, key="deposit_amt")
                else:
                    deposit_date = ""
                    deposit_amt = ""

                # ───── 小工具：把值標準化成可比對的字串 ─────
                norm = lambda v: "N/A" if v in (None, "", "N/A") else str(v)

                # ① 每次『租客 selector』改變時，把舊的輸入清掉
                if "last_selector" not in st.session_state or st.session_state.last_selector != sel_opt:
                    for k in ("curr_water_units", "curr_elec_units", "modify_calc"):
                        st.session_state.pop(k, None)
                    st.session_state.last_selector = sel_opt      # 記住這次選的人

                # ② 已計算費用 = True／False 時，也要同步清掉舊的計算
                if not calculate_done:
                    for k in ("curr_water_units", "curr_elec_units", "modify_calc"):
                        st.session_state.pop(k, None)

                # ③ 取得目前輸入（沒有就給 "N/A"）
                cur_w = norm(st.session_state.get("curr_water_units", "N/A"))
                cur_e = norm(st.session_state.get("curr_elec_units", "N/A"))

                # ④ 檢查試算是否仍然有效
                rc = st.session_state.get("modify_calc")
                calc_ok = (
                    calculate_done and
                    rc and
                    (selected_year, selected_month, norm(rc["inputs"][2]), norm(rc["inputs"][3])) ==
                    (selected_year, selected_month, cur_w, cur_e)
                )

                if not calc_ok and calculate_done:
                    st.warning("⚠️ 請先按『🔢 計算』計算金額，再儲存！")
                    st.stop()

                if st.form_submit_button("💾 儲存修改"):
                    tz_hk = pytz.timezone("Asia/Hong_Kong")
                    ts = datetime.now(tz_hk).strftime("%Y-%m-%d %H:%M:%S")
                    who = st.session_state.get("user_name", "unknown")
                    sheet_rentflow.update(f"F{gs_row}:W{gs_row}", [[
                        str(calculate_date) if calculate_done else "",
                        calculate_done,
                        water_elec_fee  if calculate_done  else "",
                        str(receive_date) if receive_done else "",
                        str(receive_done).upper(),
                        receive_amt if receive_done else "",
                        str(deposit_date) if deposit_done else "",
                        str(deposit_done).upper(),
                        deposit_amt if deposit_done else "",
                        curr_water_units, prev_water_units, water_fee,
                        curr_elec_units, prev_elec_units, elec_fee,
                        calculate_amt,
                        ts, who
                    ]])
                    st.success("✅ 已成功修改紀錄")
                    st.rerun()

    elif sub_mode == "🗑️ 刪除租金紀錄":
        st.subheader("🗑️ 刪除租金紀錄")
        if filtered_df.empty:
            st.info("目前尚無紀錄可刪除")
        else:
            filtered_df["選項"] = (
                filtered_df["租客姓名"] + "｜" +
                filtered_df["單位地址"] + "｜" +
                filtered_df["年度"].astype(str) + "-" + filtered_df["月份"].astype(str).str.zfill(2)
            )
            choice = st.selectbox("選擇要刪除的紀錄", filtered_df["選項"].tolist())
            idx = filtered_df[filtered_df["選項"] == choice].index[0]
            sheet_row = int(idx) + 2  # Google Sheets 的列數（從第2列開始）

            if st.button("⚠️ 確認刪除"):
                sheet_rentflow.delete_rows(sheet_row)
                st.warning(f"✅ 已刪除：{choice}")
                st.rerun()

    elif sub_mode == "📄 產生業主收據":
        # ── 判斷「房／座／室」token ──────────────────────────────
        PAT_ROOM  = re.compile(r"^(?:[A-Z]|[0-9]+)房$")     # A房, 1房
        PAT_WING  = re.compile(r"^[A-Z]座$")               # A座, B座
        PAT_FB    = re.compile(r"^[前後]座$")              # 前座, 後座
        PAT_UNIT  = re.compile(r"^(?:[A-Z]|[0-9]+)室$")     # A室, 1室

        def _is_part(token: str) -> bool:
            return any(p.match(token) for p in (PAT_ROOM, PAT_WING, PAT_FB, PAT_UNIT))

        def split_address(addr:str)->tuple[str,bool]:
            """回傳 (base_addr, is_partition)"""
            addr = re.sub(r"\s+", " ", addr.strip())
            head, sep, tail = addr.rpartition(" ")
            if sep and _is_part(tail):
                return head, True           # 劏房
            return addr, False              # 整個單位

        def get_ready_addresses(active_df:pd.DataFrame,
                                rentflow_df:pd.DataFrame)->list[str]:
            """
            active_df   = 當月『應交租』的租客   (我們前面已算出來)
            rentflow_df = 當月租金流程（filtered_df）
            回傳符合『整層全部已計算水電』的 base_addr 清單
            """
            # ① active_df 先標 base / is_room
            adf = active_df.copy()
            adf[["base","is_room"]] = adf["單位地址"].apply(
                lambda s: pd.Series(split_address(s))
            )

            #    每個 base 應該要出現幾間房
            need_cnt = (adf.groupby("base")
                        .size()
                        .rename("need"))

            # ② rentflow_df 只取「已計算水電」＝ TRUE
            rdf = rentflow_df.copy()
            rdf = rdf[rdf["已計算水電"].apply(lambda x: str(x).upper()=="TRUE")]
            rdf[["base","is_room"]] = rdf["單位地址"].apply(
                lambda s: pd.Series(split_address(s))
            )
            ready_cnt = (rdf.groupby("base").size().rename("ready"))
            check = (need_cnt.to_frame().join(ready_cnt, how="left").fillna(0).astype(int))
            ok_list = check[check["need"] == check["ready"]].index.tolist()
            return sorted(ok_list)
        
        def room_sort_key(addr: str, is_room: bool):
            """
            解析最後一個 token（A房 / 2室 / 前座 …）決定排序權重
            return  tuple(order_group, order_value)
            - order_group: 0=字母, 1=數字, 2=前/後, 3=非劏房
            - order_value: 用來在同 group 內排序
            """
            if not is_room:
                # 整層單位（6/F）排在最後，且依樓層排序
                # 把 6/F → 6 作數字排序；樓層越低先出
                try:
                    fl = int(addr.split("/")[-1].replace("F",""))
                except Exception:
                    fl = 999
                return (3, fl)

            token = addr.split()[-1]          # A房 / 1室 / 前座…
            base, suffix = token[:-1], token[-1]   # 去掉最後一字 (房/室/座)

            if suffix in ("房", "室"):
                if base.isalpha():            # A房 / B室
                    return (0, base.upper())
                if base.isdigit():            # 1房 / 12室
                    return (1, int(base))
            if token in ("前座", "後座"):
                return (2, 0 if token == "前座" else 1)

            # 其他情況放最後
            return (3, token)

        def generate_owner_receipt(df_month: pd.DataFrame,
                                base: str,
                                tenant_df: pd.DataFrame) -> BytesIO:
            """依樓層(base) 產生詳細業主收據 Word 並回傳 BytesIO"""

            # ===== 小工具 =====
            def _num(v):
                try:
                    if v is None:
                        return 0.0
                    s = str(v).strip()
                    if s.upper() in ("N/A", ""):
                        return 0.0
                    return float(s)
                except Exception:
                    return 0.0
            def _nz(v):
                return 0.0 if str(v).upper() in ("N/A", "", "NONE") else float(v)

            # 建立「(姓名, 地址) → 租客收費資料」查詢表
            fee_cols = ["租客姓名", "單位地址",
                        "每月固定租金", "每度水費", "固定水費",
                        "每度電費", "固定電費", "收租費"]
            fee_map = (tenant_df[fee_cols]
                    .set_index(["租客姓名", "單位地址"])
                    .applymap(_nz)      # 先把 N/A 轉 0
                    .to_dict("index"))

            # ===== Word =====
            doc = Document()
            doc.styles['Normal'].font.size = Pt(12)  # 設定字體大小
            doc.styles['Normal'].font.name = 'PMingLiU'  # 設定字體
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')

            p_title = doc.add_paragraph()
            p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_title = p_title.add_run("業主租金及水電收據")
            run_title.font.size = Pt(16)          # 標題字體大小
            run_title.font.bold = False           # 不用粗體
            run_title.font.name = 'PMingLiU'      # 指定中文字型
            run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'PMingLiU')

            p = doc.add_paragraph()
            p.add_run(f"地址：{base}\n")
            p.add_run(f"月份：{selected_year} 年 {selected_month} 月")
            
            grand_total = 0                 # ② 累加器
            mgmt_total  = 0
            parts = []

            subset = df_month[df_month["base"] == base].copy()

            subset["__sort_key"] = subset.apply(
                lambda row: room_sort_key(row["單位地址"], row["is_room"]),
                axis=1
            )

            subset = subset.sort_values("__sort_key")

            # 按租客逐一列明
            for _, r in subset.iterrows():
                key = (r["租客姓名"], r["單位地址"])
                fee = fee_map.get(key, {})
                rent  = int(fee.get("每月固定租金", 0))

                # －－ 水費 －－
                wu_curr = int(_num(r.get("本月水錶度數", 0)))
                wu_prev = int(_num(r.get("上月水錶度數", 0)))
                water_units = max(0, wu_curr - wu_prev)
                water_rate  = fee.get("每度水費", 0)
                water_fixed = fee.get("固定水費", 0)

                if water_rate > 0:          # ↖ 每度
                    water_mode = "per_unit"
                elif water_fixed > 0:       # ↖ 固定
                    water_mode = "fixed"
                else:                       # ↖ N/A → 不代收
                    water_mode = "none"

                if water_rate > 0:
                    water_mode = "per_unit"
                    value = Decimal(water_units) * Decimal(water_rate)
                    water_fee = int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
                elif water_fixed > 0:
                    water_mode = "fixed"
                    water_fee  = water_fixed
                else:
                    water_fee = 0
                    water_mode = "none"

                # －－ 電費 －－
                eu_curr = int(_num(r.get("本月電錶度數", 0)))
                eu_prev = int(_num(r.get("上月電錶度數", 0)))
                elec_units = max(0, eu_curr - eu_prev)
                elec_rate  = fee.get("每度電費", 0)
                elec_fixed = fee.get("固定電費", 0)

                if elec_rate > 0:
                    elec_mode = "per_unit"
                    value = Decimal(elec_units) * Decimal(elec_rate)
                    elec_fee  = int(value.quantize(Decimal("1"), rounding=ROUND_HALF_UP))
                elif elec_fixed > 0:
                    elec_mode = "fixed"
                    elec_fee  = elec_fixed
                else:
                    elec_fee = 0
                    elec_mode = "none"

                total = int(rent + water_fee + elec_fee)
                room_label = r["單位地址"].split()[-1] if r["is_room"] else r["單位地址"].split("/")[-1]
                parts.append(f"{room_label}:{total:.0f}")
                mgmt_fee = int(fee.get("收租費", 0))       # ← 如果 N/A 已在 _nz 變 0
                grand_total += total
                mgmt_total  += mgmt_fee

                # －－ 輸出到 Word －－
                p = doc.add_paragraph()
                p.add_run(f"租客名稱：{r['租客姓名']}\n")
                p.add_run(f"租客地址：{r['單位地址']}\n")
                if water_mode == "per_unit":
                    p.add_run(f"本月水錶度數：{wu_curr}         上月水錶度數：{wu_prev}         每度水費：{water_rate}\n")
                    p.add_run(f"水費計算： ({wu_curr}-{wu_prev}) = {wu_curr-wu_prev} × {water_rate} = {water_fee:.0f}\n")
                elif water_mode == "fixed":
                    p.add_run(f"水費： {water_fee:.0f}\n")

                if elec_mode == "per_unit":
                    p.add_run(f"本月電錶度數：{eu_curr}         上月電錶度數：{eu_prev}         每度電費：{elec_rate}\n")
                    p.add_run(f"電費計算： ({eu_curr}-{eu_prev}) = {eu_curr-eu_prev} × {elec_rate} = {elec_fee:.0f}\n")
                elif elec_mode == "fixed":
                    p.add_run(f"電費： {elec_fee:.0f}\n")

                p.add_run(
                    f"總共租金: {rent} + {water_fee} + {elec_fee} = {total}\n"
                )

            p_sum = doc.add_paragraph()
            expr = " + ".join([s.split(":")[1] for s in parts])
            p_sum.add_run(f"本層租金＋水電合計：{expr} = {grand_total:.0f}\n")
            p_sum.add_run(f"收租費合計：{mgmt_total:.0f}\n")
            p_sum.add_run(f"淨實收金額：{grand_total} - {mgmt_total} = {grand_total-mgmt_total:.0f}")

            # ===== 回傳 BytesIO =====
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf

        st.subheader("📄 產生業主收據")

        if filtered_df.empty:
            st.info(f"目前沒有 {selected_year} 年 {selected_month} 月的租金紀錄")

        # 利用工具函式挑可出單的地址
        addr_opts = get_ready_addresses(active_df, filtered_df)
        if not addr_opts:
            st.info("⚠️ 仍有房間未計算水電，暫不能產生收據")
            st.stop()

        sel_base = st.selectbox("選擇地址", addr_opts)

        if st.button("🚀 生成收據 Word"):
            # 先把 base / is_room 欄位補進 DataFrame（後續函式要用）
            filtered_df[["base","is_room"]] = filtered_df["單位地址"].apply(lambda s: pd.Series(split_address(s)))
            buf = generate_owner_receipt(filtered_df, sel_base, tenant_df)
            sel_base = sel_base.replace("/", "")
            fname = f"{selected_year}年{selected_month}月{sel_base}業主收據.docx"
            st.download_button("⬇️ 下載收據", data=buf.getvalue(), file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif main_mode == "🏢 租賃盤源管理":
    st.markdown("### 🔍 查詢間隔類型的盤源")
    layout_options = sorted(listing_df["間隔"].dropna().unique())
    layout_options.insert(0, "所有類型")  # 插入「全部」在最前面
    layout_selected = st.selectbox("📐 選擇間隔類型", layout_options)

    if layout_selected == "所有類型":
        filtered_listing = listing_df
    else:
        filtered_listing = listing_df[listing_df["間隔"] == layout_selected]
    st.write(f"共找到{len(filtered_listing)}個{layout_selected}盤源")
    st.markdown(f"### 🏢 {layout_selected}盤源一覽")
    st.data_editor(filtered_listing.set_index(pd.RangeIndex(start=1, stop=len(filtered_listing)+1)), use_container_width=True, disabled=True)
    
    sub_mode = st.radio("📋 盤源操作選項", ["➕ 新增盤源", "✏️ 更改盤源", "🗑️ 刪除盤源"], horizontal=True)
    if sub_mode == "➕ 新增盤源":
        with st.form("add_listing_form"):
            address   = st.text_input("🏠 物業地址")
            unit_type = st.selectbox("🏢 單位類型", ["劏房", "分契樓", "獨立單位"])
            layout    = st.selectbox("🛏️ 間隔", ["開放式", "套房", "一房一廳", "兩房一廳", "三房一廳"])
            gross     = st.number_input("📏 建築面積 (呎)", min_value=0.0)
            rent_amt  = st.number_input("💰 租金要求 (HKD)", min_value=0.0)
            bld_type  = st.selectbox("🏗️ 物業類型", ["唐樓", "居屋", "洋樓"])
            src_type  = st.selectbox("📌 盤源權限", ["獨家", "合作", "自己盤"])
            owner     = st.text_input("👤 業主姓名")
            owner_tel = st.text_input("📱 業主電話")
            nation    = st.selectbox("🌏 預期租客國籍", ["中國", "無限制", "外國"])
            max_occ   = st.number_input("👥 最多人數限制", min_value=1, step=1, value=1)
            remark    = st.text_area("📝 備註")
            date      = st.date_input("📅 上架日期", value=pd.Timestamp.now().date())

            if st.form_submit_button("✅ 新增"):
                dup = listing_df[(listing_df["物業地址"] == address.strip())]
                if not dup.empty:
                    st.warning("⚠️ 此地址已存在盤源，請確認是否重覆。")
                    st.stop()
                else:
                    tz_hk = pytz.timezone("Asia/Hong_Kong")
                    ts = datetime.now(tz_hk).strftime("%Y-%m-%d %H:%M:%S")
                    who = st.session_state.get("user_name", "unknown")
                    sheet_listings.append_row([
                        address, unit_type, layout, gross, rent_amt, bld_type,
                        src_type, owner, owner_tel, nation, 
                        max_occ, remark, date.strftime("%Y-%m-%d"), ts, who
                    ], value_input_option="RAW")
                    st.success("✅ 盤源已新增")
                    st.rerun()

    # ─────────────── ➋ 更改盤源 ───────────────
    elif sub_mode == "✏️ 更改盤源":
        if listing_df.empty:
            st.info("目前沒有盤源可修改。")
        else:
            selector = listing_df["物業地址"]
            choice   = st.selectbox("選擇盤源", selector)
            idx      = selector.tolist().index(choice)
            sheet_row = idx + 2
            row      = listing_df.iloc[idx]

            with st.form("edit_listing_form"):
                address   = st.text_input("🏠 物業地址", row["物業地址"])
                unit_type = st.selectbox("🏢 單位類型", ["劏房", "分契樓", "獨立單位"], index=["劏房", "分契樓", "獨立單位"].index(row["單位類型"]))
                layout    = st.selectbox("🛏️ 間隔", ["開放式", "套房", "一房一廳", "兩房一廳", "三房一廳"], index=["開放式", "套房", "一房一廳", "兩房一廳", "三房一廳"].index(row["間隔"]))
                gross     = st.number_input("📏 建築面積 (呎)", min_value=0.0, value=float(row["建築面積(呎)"]))
                rent_amt  = st.number_input("💰 租金要求 (HKD)", min_value=0.0, value=float(row["租金要求"]))
                bld_type  = st.selectbox("🏗️ 物業類型", ["唐樓", "居屋", "洋樓"],
                                         index=["唐樓", "居屋", "洋樓"].index(row["物業類型"]))
                src_type  = st.selectbox("📌 盤源權限", ["獨家", "合作", "自己盤"],
                                         index=["獨家", "合作", "自己盤"].index(row["盤源權限"]))
                owner     = st.text_input("👤 業主姓名", row["業主姓名"])
                owner_tel = st.text_input("📱 業主電話", row["業主電話"])
                nation    = st.selectbox("🌏 預期租客國籍", ["中國", "無限制", "外國"],
                                         index=["中國", "無限制", "外國"].index(row["預期租客國籍"]))
                max_occ   = st.number_input("👥 最多人數限制", min_value=1, step=1, value=1 if str(row["最多人數限制"])=="N/A" else int(row["最多人數限制"]))
                remark    = st.text_area("📝 備註", row["備註"])

                if st.form_submit_button("💾 儲存修改"):
                    tz_hk = pytz.timezone("Asia/Hong_Kong")
                    ts = datetime.now(tz_hk).strftime("%Y-%m-%d %H:%M:%S")
                    who = st.session_state.get("user_name", "unknown")
                    sheet_listings.update(
                        f"A{sheet_row}:O{sheet_row}",
                        [[address, unit_type, layout, gross, rent_amt, bld_type, 
                          src_type, owner, owner_tel, nation,
                          max_occ, remark, str(row["上架日期"]), ts, who]], value_input_option="RAW"
                    )
                    st.success("✅ 已成功更改盤源")
                    st.rerun()

    # ─────────────── ➌ 刪除盤源 ───────────────
    elif sub_mode == "🗑️ 刪除盤源":
        if listing_df.empty:
            st.info("目前沒有盤源可刪除。")
        else:
            selector = listing_df["物業地址"]
            choice   = st.selectbox("選擇盤源", selector)
            idx      = selector.tolist().index(choice)
            sheet_row = idx + 2
            if st.button("⚠️ 確認刪除"):
                sheet_listings.delete_rows(sheet_row)
                st.warning(f"已刪除：{choice}")
                st.rerun()